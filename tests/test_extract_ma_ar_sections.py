import os
import shutil
import pytest
from docx import Document
from backend.extract.ma_ar_parser import extract_ma_ar_sections
from tests.utils.test_helpers import save_sections_to_docx

TEST_DOCX = "tests/SectionI&II-fortest.docx"
OUTPUT_DIR = "tests/tmp_output"

@pytest.fixture(scope="module")
def test_docx_file():
    """Create a synthetic test Word doc for MA & AR."""
    doc = Document()
    doc.add_paragraph("前言内容")
    doc.add_paragraph("管理层认定")
    doc.add_paragraph("这是管理层的第一段。")
    doc.add_paragraph("这是管理层的第二段。")
    doc.add_paragraph("审计师报告")
    doc.add_paragraph("这是审计师报告的第一段。")
    doc.save(TEST_DOCX)
    yield TEST_DOCX
    os.remove(TEST_DOCX)

@pytest.fixture
def clean_output_dir():
    if os.path.exists(OUTPUT_DIR):
        shutil.rmtree(OUTPUT_DIR)
    os.makedirs(OUTPUT_DIR)
    yield OUTPUT_DIR
    shutil.rmtree(OUTPUT_DIR)

def test_extract_sections(test_docx_file):
    doc = Document(test_docx_file)
    ma, ar = extract_ma_ar_sections(doc)

    assert isinstance(ma, list)
    assert isinstance(ar, list)
    assert "管理层的第一段" in "".join(ma)
    assert "审计师报告的第一段" in "".join(ar)

def test_output_docx_files(test_docx_file, clean_output_dir):
    doc = Document(test_docx_file)
    ma, ar = extract_ma_ar_sections(doc)
    ma_path, ar_path = save_sections_to_docx(ma, ar, clean_output_dir)

    assert os.path.exists(ma_path)
    assert os.path.exists(ar_path)

    ma_doc = Document(ma_path)
    ar_doc = Document(ar_path)

    assert "管理层的第一段" in "\n".join(p.text for p in ma_doc.paragraphs)
    assert "审计师报告的第一段" in "\n".join(p.text for p in ar_doc.paragraphs)
