from docx import Document
import os

def save_sections_to_docx(ma_text, ar_text, output_dir):
    """Test-only helper to save MA/AR content into docx for inspection."""
    os.makedirs(output_dir, exist_ok=True)

    ma_doc = Document()
    ma_doc.add_heading("管理层认定 (MA Section)", level=1)
    for para in ma_text:
        ma_doc.add_paragraph(para)
    ma_path = os.path.join(output_dir, "MA_section.docx")
    ma_doc.save(ma_path)

    ar_doc = Document()
    ar_doc.add_heading("审计师报告 (AR Section)", level=1)
    for para in ar_text:
        ar_doc.add_paragraph(para)
    ar_path = os.path.join(output_dir, "AR_section.docx")
    ar_doc.save(ar_path)

    return ma_path, ar_path
