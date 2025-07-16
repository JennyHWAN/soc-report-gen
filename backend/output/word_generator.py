import os
import tempfile
import shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from backend.extract.ma_ar_parser import extract_ma_ar_sections


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_paragraphs(doc, paragraphs, indent=False):
    for p in paragraphs:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(p.strip())
        if indent:
            paragraph.paragraph_format.first_line_indent = Pt(28.3)  # ≈1cm
        run.font.size = Pt(12)


def add_signature_block(doc, signer_name, lines_before=4):
    for _ in range(lines_before):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = 2  # Right-aligned
    run = p.add_run(signer_name.strip())
    run.font.size = Pt(12)


def generate_ma_ar_docx(tmp_word_path, word_file, output_base_path: str) -> str:
    """
    Generates a DOCX report (Part I & II) from an uploaded Word file.

    Args:
        tmp_word_path: The word uploaded.
        word_file: A file-like object containing the .docx Word input.
        output_base_path: Path without extension, e.g., 'generated_reports/Part_I_II'

    Returns:
        The full path to the generated DOCX file.
    """
    doc = Document(tmp_word_path)

    # Extract raw text
    ma_text, ar_text = extract_ma_ar_sections(doc)

    # === Build New Word Document ===
    output_doc = Document()

    # Set default font
    style = output_doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")

    # Part I – 管理层认定
    add_section_heading(output_doc, "第一部分 – 管理层认定")
    add_paragraphs(output_doc, ma_text[:-1])
    add_signature_block(output_doc, signer_name=ma_text[-1], lines_before=4)

    output_doc.add_page_break()

    # Part II – 审计师报告
    add_section_heading(output_doc, "第二部分 – 独立服务审计师报告")
    add_paragraphs(output_doc, ar_text[:-3])
    add_signature_block(output_doc, signer_name="\n".join(ar_text[-3:]), lines_before=8)

    # Save output
    final_docx_path = output_base_path + ".docx"
    output_doc.save(final_docx_path)

    return final_docx_path
