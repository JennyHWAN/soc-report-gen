from pylatexenc.latexencode import unicode_to_latex
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import subprocess
import streamlit as st

def format_as_latex(text_list):
    # Escape LaTeX special characters and join paragraphs
    # Only escape actual LaTeX special characters — leave Chinese characters untouched
    def escape_latex(text):
        return (text.replace('\\', r'\textbackslash{}')
                .replace('&', r'\&')
                .replace('%', r'\%')
                .replace('$', r'\$')
                .replace('#', r'\#')
                .replace('_', r'\_')
                .replace('{', r'\{')
                .replace('}', r'\}')
                .replace('~', r'\textasciitilde{}')
                .replace('^', r'\textasciicircum{}'))

    return "\n\n".join(escape_latex(p) for p in text_list)

def format_paragraphs_to_latex(paragraphs, indent=False):
    """
    Safely include paragraphs in LaTeX using raw Unicode (for XeLaTeX).
    """
    return "\n\n".join(
        (r"\par\noindent " if indent else "") + p for p in paragraphs
    )

def latex_vspace_lines(lines=4):
    """
    Returns a vertical space block for a given number of lines (approx. 1em per line).
    """
    return "\n".join([r"\vspace*{3em}"] * lines)

def latex_signature_block(entity_name, lines_before=4):
    """
    Creates a signature block with spacing and organization name.
    """
    spacing = latex_vspace_lines(lines_before)
    return f"{spacing}\n{entity_name}"

def latex_document_wrapper(body: str, font_main='TeX Gyre Termes', font_cjk='Noto Sans CJK SC') -> str:
    """
    Wraps LaTeX body content in a complete XeLaTeX document structure.
    """
    return rf"""
    \documentclass[12pt]{{article}}
    \usepackage[margin=1in]{{geometry}}
    \usepackage{{xeCJK}}
    \usepackage{{fontspec}}
    \setmainfont{{{font_main}}}
    \setCJKmainfont{{{font_cjk}}}
    \clubpenalty=10000
    \widowpenalty=10000
    
    \begin{{document}}
    
    {body}
    
    \end{{document}}
    """

def render_latex_to_pdf(tex_path):
    output_dir = os.path.dirname(tex_path)
    os.makedirs(output_dir, exist_ok=True)
    try:
        subprocess.run(["xelatex", "-output-directory", output_dir, tex_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except subprocess.CalledProcessError as e:
        st.error("❌ XeLaTeX compilation failed.")
        st.text(e.stderr.decode("utf-8"))
        raise
    except subprocess.TimeoutExpired:
        st.error("❌ XeLaTeX timed out.")
        raise
    return tex_path.replace(".tex", ".pdf")


# TODO: split the word generation and pdf generation function into two to make it more structured
# This is for generating docx file
def generate_docx_with_body(body: str, output_path=os.path.join(os.getcwd(), "generated_reports/Part_I_II.docx")):
    doc = Document()

    # Apply font globally (not perfect, needs to be applied per paragraph or style)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'  # Substitute for 'TeX Gyre Termes'
    font.size = Pt(12)

    # For CJK support — optional and needs MS fonts
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')

    # Add body content
    for para in body.split("\n\n"):
        doc.add_paragraph(para.strip())

    # Add custom text before page break
    doc.add_paragraph("here are some text.")

    # Page break
    doc.add_page_break()

    # Text after page break
    doc.add_paragraph("test for new page.")

    # Save the document
    doc.save(output_path)
    return output_path
